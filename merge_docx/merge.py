from docx import Document
import docx
import os

from .utils.handle_floats import handle_floats
from .utils.handle_hyperlinks import handle_hyperlinks
from .utils.handle_styles import handle_styles
from .utils.handle_inlines import handle_inlines
from .utils.handle_sections import handle_sections
from .utils.handle_headers_footers import handle_headers_footers
from .utils.handle_footnotes import handle_footnotes
from .utils.handle_numbers import handle_numbers

from pkg_resources import resource_filename
BLANK_FILE = resource_filename(__name__, 'blank.docx')

# The first merge will be into a blank .docx file containing a footnote part.
# This allows subsequently merged documents to preserve their footnotes.
def blank_merge(file):
    file_no_elem = file + 'no_elem.docx'
    # Remove hyperlinks
    handle_hyperlinks(file, file_no_elem)

    # blank.docx contains the necessary footnote part, so this is the template
    merged_document = Document(BLANK_FILE)

    # The document to be merged in exists as the second file in the list. Note
    # that this script only merges 2 files at a time.
    sub_doc = Document(file_no_elem)

    # Add the footnotes from sub_doc to merged_document.
    handle_footnotes(merged_document, sub_doc)

    # Add each style in sub_styles to merged_styles.
    handle_styles(merged_document, sub_doc)

    # Remove the temporary files generated without floating elements. If a file
    # is being merged into itself, there will only be one document generated 
    # without floats/hyperlinks. Trying to delete both documents will throw an exception,
    # but I don't care as long as the file gets deleted. 
    try:
        os.remove(file_no_elem)
    except OSError:
        pass

    # Add the inline images from sub_doc into the merged document
    handle_inlines(merged_document, sub_doc)



    # Preserve floating images
    handle_floats(merged_document, sub_doc)


    

    # Add numbering elements from sub_doc into the merged document
    handle_numbers(merged_document, sub_doc)

    # Merge the document bodies   
    for element in sub_doc.element.body:
        merged_document.element.body.append(element)

    # Regroup end-sections lost in the merge.
    handle_sections(merged_document)

    # Remove all header and footer references.
    handle_headers_footers(merged_document)

    # Save the altered file
    destination = 'temp.docx'

    merged_document.save(destination)

    # Return the file location to be used in the next merge
    return destination



# Merge 'file' into 'template'
def merge_docx(template, file, destination):

    template = blank_merge(template)

    # Save with a .docx file extension
    if destination[-5:] != '.docx':
        destination += '.docx'

    template_no_elem = template + 'no_elem.docx'
    file_no_elem = file + 'no_elem.docx'

    # Remove hyperlinks
    handle_hyperlinks(template, template_no_elem)
    handle_hyperlinks(file, file_no_elem)

    # Template was overwritten as 'temp.docx' in blank_merge, but is not needed
    # at this point because it has already been opened and processed
    os.remove(template)

    # Use the first document as a template file. All formatting will be 
    # inherited from this file. All media and relationships from this file are
    # preserved.
    merged_document = Document(template_no_elem)

    # Remove the blank paragraph created by blank_merge
    blank_para = merged_document.paragraphs[0]
    p = blank_para._element
    p.getparent().remove(p)
    p._p = p._element = None

    # The document to be merged in exists as the second file in the list. Note
    # that this script only merges 2 files at a time.
    sub_doc = Document(file_no_elem)

    # Add the footnotes from sub_doc to merged_document.
    handle_footnotes(merged_document, sub_doc)

    # Add each style in sub_styles to merged_styles.
    handle_styles(merged_document, sub_doc)

    # Remove the temporary files generated without floating elements. If a file
    # is being merged into itself, there will only be one document generated 
    # without floats/hyperlinks. Trying to delete both documents will throw an exception,
    # but I don't care as long as the file gets deleted. 
    try:
        os.remove(template_no_elem)
        os.remove(file_no_elem)
    except OSError:
        pass

    # Add the inline images from sub_doc into the merged document
    handle_inlines(merged_document, sub_doc)


    # Add numbering elements from sub_doc into the merged document
    handle_numbers(merged_document, sub_doc)


    # Merge the document bodies   
    for element in sub_doc.element.body:
        merged_document.element.body.append(element)

    # Regroup end-sections lost in the merge.
    handle_sections(merged_document)

    # Remove all header and footer references.
    handle_headers_footers(merged_document)

    # Save the altered file
    merged_document.save(destination)




    